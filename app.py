import pandas as pd
import numpy as np
import os
import re
from datetime import datetime
from docx import Document, shared
#import comtypes.client
import time
import warnings
import streamlit as st
import pypandoc
import subprocess
#from io import BytesIO

warnings.filterwarnings('ignore', category=UserWarning, module='openpyxl')


def initialize_doc(template_file_path: str) -> Document:
    """
    Initializes a word document with the given path.
    """
    print('Initializing document...')
    try:
        doc = Document(template_file_path)
        # Setting the Normal font works:
        font = doc.styles['Normal'].font
        font.name = 'Calibri'
        font.size = shared.Pt(9)
    except Exception as e:
        raise Exception(f"Error initializing the document. {e}")
    
    print("Document initialized succesfully")
    
    return doc

def input_files_structure(df: pd.DataFrame, type: str) -> bool:
    """
    Checks if the input files structure is correct.
    """
    print(f'Check file Structure for {type}...')
    if type == 'annexes':
        df_structure = pd.DataFrame({'index': ['seq', 'row_number', 'reg', 'requisito', 'item', 
                                               'medio_de_cumplimiento', 'si', 'na', 'docx_table', 'docx_row','docx_cump_col', 'docx_si_col', 'docx_na_col'], 
                                     0: ['int64']*2 + ['object']*6 + ['float64']*5})
        if not df_structure.equals(df.dtypes.reset_index()):
            st.error("The annexes file structure is not correct.")
            return False
        elif df['medio_de_cumplimiento'].str.contains('CHECK WITH SYSTEMS.CAMO').any():
            st.error("The annexes file structure is not correct. CHECK WITH SYSTEMS.CAMO is not allowed.")
            return False
    elif type == 'equipment':
        df_structure = pd.DataFrame({'index': ['reg', 'ac_model', 'msn', 'mfg_date', 'equipment', 'manufacturer', 'partno', 'mel_item', 'serialno', 'fin',
                        'date_inst', 'operation', 'fabricante', 'alt_max', 'mode_s_code'],
                        0: ['object', 'object', 'int64', 'datetime64[ns]', 'object', 'object', 'object', 'object', 'object', 'object', 'datetime64[ns]', 'object', 'object', 'object', 'object'] })
        if not df_structure.equals(df.dtypes.reset_index()):
            return False

    print("Input files structure checked succesfully")

    return True

def upload_annexes_dataframe_to_array(annexes_file_path: str) -> np.ndarray:
    """
    Uploads the annexes file to a numpy array.
    """
    print('Uploading annexes file..')
    try:
        annexes_df = pd.read_excel(annexes_file_path)
        annexes_df.columns = [col.lower() for col in annexes_df.columns]
        if not input_files_structure(annexes_df, 'annexes'):
            raise Exception("The annexes file structure is not correct.")
        else:
            annexes_df = annexes_df.dropna(subset=['docx_table'])\
                                    .replace(np.nan, None)\
                                    .fillna('')
            annexes_df[[col for col in annexes_df.columns if re.search('docx', col, re.IGNORECASE) is not None]] \
                = annexes_df[[col for col in annexes_df.columns if re.search('docx', col, re.IGNORECASE) is not None]].astype('int')
            
            annexes_array = annexes_df.to_numpy()
    except Exception as e:
        raise Exception(f"Error uploading the annexes file. {e}")

    print("Annexes file uploaded succesfully")

    return annexes_array

def upload_equipment_dataframe_to_array(equipment_file_path: str) -> np.ndarray:
    """
    Uploads the equipment file to 4 numpy arrays.
    """
    print('Uploading equipment file..')
    try:
        equipment_df = pd.read_excel(equipment_file_path)
        equipment_df.columns = [col.lower() for col in equipment_df.columns]
        if not input_files_structure(equipment_df, 'equipment'):
            raise Exception("The equipment file structure is not correct.")
        else:
            equipment_df = equipment_df.replace(np.nan, None)\
                                        .fillna('N/A')
            
            array_lvo = equipment_df.loc[equipment_df['operation'] == 'LVO', ['equipment', 'manufacturer', 'partno', 'mel_item']].to_numpy()
            array_rvsm = equipment_df.loc[equipment_df['operation'] == 'RVSM', ['equipment', 'manufacturer', 'partno', 'mel_item']].to_numpy()
            array_pbn = equipment_df.loc[equipment_df['operation'] == 'PBN', ['equipment', 'manufacturer', 'partno', 'mel_item']].to_numpy()
            array_cpdlc = equipment_df.loc[equipment_df['operation'] == 'CPDLC', ['equipment', 'manufacturer', 'partno', 'mel_item']].to_numpy()

    except Exception as e:
        raise Exception(f"Error uploading the equipment file. {e}")

    print("Equipment file uploaded succesfully")

    return (array_lvo, array_rvsm, array_pbn, array_cpdlc)

def populate_annexes(document: Document, array: np.ndarray) -> Document:
    """
    Populates the annexes in the word document.
    """
    print('Populating annnexes field..')
    try:
        for array_row, docx_row in enumerate(array[:,9]):
            if array[array_row, 8] == 4:
                # medio de cumplimiento
                document.tables[array[array_row, 8]].cell(docx_row, 6).text = array[array_row, 5]
                # si
                document.tables[array[array_row, 8]].cell(docx_row, 10).text = array[array_row, 6]
                # na
                document.tables[array[array_row, 8]].cell(docx_row, 12).text = array[array_row, 7]

            elif array[array_row, 8] == 5:
                # medio de cumplimiento
                document.tables[array[array_row, 8]].cell(docx_row, 8).text = array[array_row, 5]
                # si
                document.tables[array[array_row, 8]].cell(docx_row, 9).text = array[array_row, 6]
                # na
                document.tables[array[array_row, 8]].cell(docx_row, 10).text = array[array_row, 7]
            elif array[array_row, 8] == 6:
                # medio de cumplimiento
                document.tables[array[array_row, 8]].cell(docx_row, 2).text = array[array_row, 5]
                # si
                document.tables[array[array_row, 8]].cell(docx_row, 3).text = array[array_row, 6]
                # na
                document.tables[array[array_row, 8]].cell(docx_row, 4).text = array[array_row, 7]
            else:
                raise ValueError('Invalid value in column 8')
        
        # specific case for CAT.IDE.A.125 where a row must be duplicated
        for i in range(17, 24):
            # medio de cumplimiento
            document.tables[4].cell(i, 6).text = 'No están aprobadas las operaciones VFR para aviones de Transporte de Pasajeros'
            # na
            document.tables[4].cell(i, 12).text = 'X'


        # specific case for Datalink Reg 29/2009 
        # medio de cumplimiento
        document.tables[4].cell(213, 6).text = 'Cubierto por el Diseño de Tipo de la aeronave. Equipos instalados se describen en Anexo E'
        # si
        document.tables[4].cell(213, 10).text = 'X'


        # specific case for CAT.IDE.A.355
        # medio de cumplimiento
        document.tables[4].cell(220, 6).text = 'Campo no especificado'
        # na
        document.tables[4].cell(220, 12).text = 'X'

        # specific case for CAT.OP.MPA.126
        # medio de cumplimiento
        document.tables[4].cell(236, 6).text = 'Cubierto por el Diseño de Tipo de la aeronave. Se describe en AFM LIM-22-FMS'
        # si
        document.tables[4].cell(236, 10).text = 'X'


        # specific case for 26.250
        # medio de cumplimiento
        document.tables[7].cell(3, 2).text = 'El sistema CDLS permite a la tripulación de cabina de pasajeros solicitar acceso al compartimiento de vuelo mediante un código de seguridad'
        # si
        document.tables[7].cell(3, 3).text = 'X'
    except Exception as e:
        raise Exception(f"Error populating the annexes. {e}")

    print("Annexes populated succesfully")
    return document

def populate_equipment(document: Document, arrays: tuple[np.ndarray]) -> Document:
    """
    Populates the equipment tables in the word document.
    """
    print('Populating equipment tables..')
    array_lvo, array_rvsm, array_pbn, array_cpdlc = arrays
    table_dict = {'LVO': (12, array_lvo), 'RVSM': (14, array_rvsm), 'PBN': (22, array_pbn), 'CPDLC': (31, array_cpdlc)}

    for dic in table_dict:
        try:
            table_num = table_dict[dic][0]
            array = table_dict[dic][1]
            table_len = len(document.tables[table_num].rows)
            array_len = array.shape[0]
            for array_row_num, array_row_value in enumerate(array):
                for array_col_num, array_col_value in enumerate(array_row_value):
                    document.tables[table_num].cell(1 + array_row_num, array_col_num).text = array_col_value
            
            # delete spare rows if any
            if table_len > array_len:
                for table_row in range (table_len -1 , array_len, -1):
                    document.tables[table_num]._tbl.remove(document.tables[table_num].rows[table_row]._tr)
        except Exception as e:
            print(f'Error in table {dic}: {e}')
            continue

    print("Equipment populated succesfully")
    return document

def populate_rest_of_data(document: Document, equipment_file_path: str) -> Document:
    """
    Populates the rest of the data in the word document.
    """
    print('Populating rest of data in the form..')
    try:
        equipment_df = pd.read_excel(equipment_file_path)
        equipment_df.columns = [col.lower() for col in equipment_df.columns]
        if not input_files_structure(equipment_df, 'equipment'):
            raise Exception("The equipment file structure is not correct.")
        else:
            equipment_df = equipment_df.replace(np.nan, None)\
                                        .fillna('N/A')

            ################# populate Datos de la aeronave
            # Fabricante
            document.tables[0].cell(2, 3).text = equipment_df['fabricante'].unique()[0]
            # Tipo/Variante
            document.tables[0].cell(2, 14).text = equipment_df['ac_model'].unique()[0]
            # Matricula
            document.tables[0].cell(3, 3).text = equipment_df['reg'].unique()[0]
            # MSN Noº
            document.tables[0].cell(3, 10).text = str(equipment_df['msn'].unique()[0])
            # CofA
            document.tables[0].cell(7, 6).text = equipment_df['mfg_date'].unique()[0].strftime('%d/%b/%Y')

            ################# populate transponder data
            #Marca/Tipo/Variante
            document.tables[15].cell(1, 1).text = equipment_df.loc[equipment_df['fin'] == '1SH1', 'manufacturer'].item()
            # P/N
            document.tables[15].cell(1, 2).text = equipment_df.loc[equipment_df['fin'] == '1SH1', 'partno'].item()
            # ICAO CODE
            document.tables[15].cell(1, 3).text = equipment_df.loc[equipment_df['fin'] == '1SH1', 'mode_s_code'].item()
    except Exception as e:
        raise Exception(f"Error populating the rest of the data. {e}")

    print("Rest of data populated succesfully")
    return document

def save_doc(document: Document, equipment_file_path: str, output_file_path: str) -> str:
    """
    Saves the word documenta and returns the resulting path.
    """
    print('Saving .docx file')
    try:
        df = pd.read_excel(equipment_file_path)
        df.columns = [col.lower() for col in df.columns]
        reg = df['reg'].unique()[0]
        specific_path = f'{output_file_path} {reg}.docx'
        document.save(specific_path)
    except Exception as e:
        raise Exception(f"Error saving the document. {e}")

    print("Document saved succesfully")
    return specific_path

def convert_doc_to_pdf(input_file_path: str, output_pdf_file_path: str) -> None:
    """
    Converts a .docx file to .pdf using pypandoc and ignores the 'No pandoc was found' error.
    """
    print("Converting .docx file to .pdf")
    try:
        pypandoc.convert_file(input_file_path, 'pdf', outputfile=output_pdf_file_path)
        print("Document converted successfully")
    except Exception as e:
        # Check for the specific "No pandoc was found" error
        if "No pandoc was found" in str(e):
            print("Warning: Pandoc is not installed. Skipping PDF conversion.")
        else:
            # Re-raise other exceptions
            raise Exception(f"Error converting the document. {e}")

    print("Document converted succesfully")

def main(template_file_path: str, annexes_file_path: str, equipment_file_path: str, output_file_path: str) -> str:
    """
    Main function. Returns the resulting files path (.xlsx and .pdf)
    """
    try:
        document = initialize_doc(template_file_path)
        annexes_array = upload_annexes_dataframe_to_array(annexes_file_path)
        document = populate_annexes(document, annexes_array)
        equipment_arrays = upload_equipment_dataframe_to_array(equipment_file_path)
        document = populate_equipment(document, equipment_arrays)
        document = populate_rest_of_data(document, equipment_file_path)
        output_docx_file_path = save_doc(document, equipment_file_path, output_file_path)
        output_pdf_file_path = output_docx_file_path.replace('.docx', '.pdf')
        convert_doc_to_pdf(output_docx_file_path, output_pdf_file_path)
    except Exception as e:
        st.error(f"Error in main function. {e}")
    
    return output_docx_file_path, output_pdf_file_path

# Define a dictionary of usernames and passwords
USER_CREDENTIALS = {
    "admin": "camo",
    "user1": "Vu3l1ng2024!",
    # Add more users as needed
}

# Login function
def login():
    """Display login form and check credentials."""
    st.title("Login")
    # Get username and password from the user
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    # Login button
    if st.button("Login"):
        if username in USER_CREDENTIALS and USER_CREDENTIALS[username] == password:
            st.session_state.logged_in = True
            st.session_state.username = username
            st.experimental_rerun()  # Reload to access the main application
        else:
            st.error("Invalid username or password")



# Define a dictionary of usernames and passwords
USER_CREDENTIALS = {
    "camo": {"username": "camo", "password": "Vu3l1ng2024!"}
    # Add more users as needed
}

# Login function
def login():
    """Display login form and check credentials."""
    st.title("Login")
    # Get username and password from the user
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    # Login button
    if st.button("Login"):
        # Debugging: Print the input values to check if they are being entered correctly
        print(f"Attempting login with username: {username}, password: {password}")

        # Check if username and password match
        if username in USER_CREDENTIALS and USER_CREDENTIALS[username]['password'] == password:
            st.session_state.logged_in = True
            st.session_state.username = username
            st.success(f"Logged in as {username}")  # Confirmation message
            st.rerun()  # Reload to access the main application
        else:
            st.error("Invalid username or password")


def main_app():
    """Main application logic."""

    #TODO
    try:
        result = subprocess.run(['texlive-full', '--version'], capture_output=True, text=True, check=True)
        st.write(result.stdout)
    except Exception as e:
        st.error(f"Error: {e}")
    # Static variables
    template_file = 'AC-ACAM-P01-F31.docx'
    abs_dir_path = os.getcwd()
    template_file_path = os.path.join(abs_dir_path, 'config', template_file)
    output_file_path = os.path.join(abs_dir_path, 'AC-ACAM-P01-F31 Ed.02 Declaración equipamiento AOC Avión')

    st.title("AC-ACAM-P01-F31 Generation (Single Aircraft)")

    # Step 1: Upload the annexes file
    col1, col2 = st.columns([2, 1])
    with col1:
        file1 = st.file_uploader(
            "Upload the Excel file with the AC-ACAM-P01-F31 annexes previously downloaded from AMOS (Filename not relevant)",
            type="xlsx",
            key="file1"
        )
    with col2:
        st.image("file1_instructions.png", caption="How to get AC-ACAM-P01-F31 annexes file", use_container_width=True)

    # Step 2: Upload the equipment file
    col3, col4 = st.columns([2, 1])
    with col3:
        file2 = st.file_uploader(
            "Upload the Excel file with the LVO/PBN/RVSM/CPDLC equipment previously downloaded from AMOS (Filename not relevant)",
            type="xlsx",
            key="file2"
        )
    with col4:
        st.image("file2_instructions.png", caption="How to get AC-ACAM-P01-F31 equipment file", use_container_width=True)

    # Use session state to persist uploaded files
    if file1:
        st.session_state["annexes_file"] = file1
    if file2:
        st.session_state["equipment_file"] = file2

    # Only proceed if both files are uploaded
    if "annexes_file" in st.session_state and "equipment_file" in st.session_state:
        annexes_file_path = st.session_state["annexes_file"]
        equipment_file_path = st.session_state["equipment_file"]

        # Step 3: Execute main function
        if st.button('Generate AC-ACAM-P01-F31'):
            with st.spinner('Generating AC-ACAM-P01-F31...'):
                output_docx_file_path, output_pdf_file_path = main(template_file_path, annexes_file_path, equipment_file_path, output_file_path)
                st.session_state["output_docx"] = output_docx_file_path
                st.session_state["output_pdf"] = output_pdf_file_path
                st.success('AC-ACAM-P01-F31 generated successfully!')

    # Step 4: Display download buttons and warnings if files are generated
    if "output_docx" in st.session_state and "output_pdf" in st.session_state:
        col5, col6 = st.columns(2)
        # debugging TODO borrar
        st.write(f"{os.listdir()}")
        with col5:
            with open(st.session_state["output_docx"], 'rb') as f:
                st.download_button(
                    'Download AC-ACAM-P01-F31 in .xlsx format', 
                    f, 
                    file_name=st.session_state["output_docx"]
                )

        with col6:
            with open(st.session_state["output_pdf"], 'rb') as f:
                st.download_button(
                    'Download AC-ACAM-P01-F31 in .pdf', 
                    f, 
                    file_name=st.session_state["output_pdf"]
                )

        # Step 7: Display Note/Warning with Image
        st.write("Important Note!")
        col7, col8 = st.columns([1, 2])
        with col7:
            st.image("final_note_efb_mel_mmel.png", caption="Important", use_container_width=True)
        with col8:
            st.warning(
                """Please consider checking/updating the latest version of MEL and MMEL as well as fulfilling the table equipment for EFB. 
                If you encounter issues, contact systems.camo@vueling.com."""
            )


# Entry point of the app
if __name__ == '__main__':
    start = time.time()
    # Initialize session state if it's not already initialized
    if "logged_in" not in st.session_state:
        st.session_state.logged_in = False

    # After login check, automatically show the main app
    if st.session_state.logged_in:
        main_app()  # Call main app logic if logged in
    else:
        login()  # Show login form if not logged in

    print(f'Elapsed time: {time.time() - start} seconds')